"""
Google Document AI integration for processing legal documents
"""
import json
import uuid
from typing import Dict, Any, List
from datetime import datetime
from google.cloud import documentai
from google.api_core import exceptions as gcp_exceptions
import tempfile
import os
import re

from models.document_models import DocumentAnalysis, DocumentMetadata, Clause
from utils.config import get_settings

class DocumentProcessor:
    """Processes legal documents using Google Document AI"""
    
    def __init__(self):
        self.settings = get_settings()
        
        if not self.settings.google_cloud_project_id:
            raise ValueError("Google Cloud project ID is required")
            
        # Initialize client with explicit credentials if provided
        if self.settings.google_application_credentials:
            import os
            os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = self.settings.google_application_credentials
            
        try:
            self.client = documentai.DocumentProcessorServiceClient()
        except Exception as e:
            print(f"Error initializing Document AI client: {e}")
            raise e
            
        self.project_id = self.settings.google_cloud_project_id
        self.location = self.settings.document_ai_location
        self.processor_id = self.settings.document_ai_processor_id
        
    async def process_document(self, content: bytes, filename: str, session_id: str) -> DocumentAnalysis:
        """
        Process a legal document using Document AI and extract structured data
        """
        try:
            # Configure the processor path
            processor_path = self.client.processor_path(
                self.project_id, self.location, self.processor_id
            )
            
            # If DOCX, convert to PDF for Document AI compatibility
            mime_type = self._get_mime_type(filename)
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                try:
                    content = self._convert_docx_bytes_to_pdf_bytes(content)
                    mime_type = 'application/pdf'
                except Exception as conv_err:
                    raise ValueError(
                        f"DOCX conversion failed: {conv_err}. Please upload a PDF file instead."
                    )

            # Create document object
            raw_document = documentai.RawDocument(
                content=content,
                mime_type=mime_type
            )
            
            # Create process request
            request = documentai.ProcessRequest(
                name=processor_path,
                raw_document=raw_document
            )
            
            # Process the document
            result = self.client.process_document(request=request)
            document = result.document
            
            # Extract metadata and clauses
            metadata = self._extract_metadata(document)
            clauses = self._extract_clauses(document)
            
            # Calculate confidence score
            confidence_score = self._calculate_confidence(document)
            
            return DocumentAnalysis(
                document_id=session_id,
                document_name=filename,
                metadata=metadata,
                clauses=clauses,
                processing_timestamp=datetime.utcnow(),
                confidence_score=confidence_score
            )
            
        except gcp_exceptions.GoogleAPIError as e:
            self._handle_processing_error(e, filename, session_id)
        except Exception as e:
            self._handle_processing_error(e, filename, session_id)

    def _convert_docx_bytes_to_pdf_bytes(self, docx_bytes: bytes) -> bytes:
        """Convert DOCX bytes to PDF bytes.
        Tries docx2pdf first (best fidelity). If unavailable, falls back to
        extracting text via python-docx and generating a simple PDF with reportlab.
        """
        # Try docx2pdf (requires MS Word on Windows)
        try:
            from docx2pdf import convert as docx2pdf_convert
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "input.docx")
                pdf_path = os.path.join(tmpdir, "output.pdf")
                with open(docx_path, 'wb') as f:
                    f.write(docx_bytes)
                docx2pdf_convert(docx_path, pdf_path)
                with open(pdf_path, 'rb') as f:
                    return f.read()
        except Exception:
            # Fallback: create a text-based PDF
            try:
                from docx import Document as DocxDocument
                from reportlab.lib.pagesizes import LETTER
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import inch
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                import textwrap

                with tempfile.TemporaryDirectory() as tmpdir:
                    # Extract text from DOCX
                    docx_path = os.path.join(tmpdir, "input.docx")
                    with open(docx_path, 'wb') as f:
                        f.write(docx_bytes)
                    doc = DocxDocument(docx_path)
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

                    # Generate PDF
                    pdf_path = os.path.join(tmpdir, "output.pdf")
                    c = canvas.Canvas(pdf_path, pagesize=LETTER)
                    width, height = LETTER

                    # Register a basic font for better unicode handling if available
                    try:
                        pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                        font_name = 'DejaVuSans'
                    except Exception:
                        font_name = 'Helvetica'

                    c.setFont(font_name, 11)
                    left_margin = 1 * inch
                    right_margin = 1 * inch
                    top_margin = 1 * inch
                    bottom_margin = 1 * inch
                    max_width = width - left_margin - right_margin
                    line_height = 14
                    y = height - top_margin

                    wrapper = textwrap.TextWrapper(width=100)

                    for para in paragraphs:
                        # Wrap paragraph to fit page width roughly
                        lines = wrapper.wrap(para)
                        for line in lines:
                            if y <= bottom_margin:
                                c.showPage()
                                c.setFont(font_name, 11)
                                y = height - top_margin
                            c.drawString(left_margin, y, line)
                            y -= line_height
                        # Add blank line between paragraphs
                        y -= line_height // 2

                    c.save()
                    with open(pdf_path, 'rb') as f:
                        return f.read()
            except Exception as fallback_err:
                raise RuntimeError(
                    "DOCX conversion not available. Please upload a PDF file instead."
                ) from fallback_err
    
    def _get_mime_type(self, filename: str) -> str:
        """Get MIME type from filename"""
        if filename.lower().endswith('.pdf'):
            return 'application/pdf'
        elif filename.lower().endswith('.docx'):
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return 'application/pdf'  # Default
    
    def _extract_metadata(self, document: Any) -> DocumentMetadata:
        """Extract metadata from processed document"""
        text = getattr(document, 'text', '') or ''
        entities = document.entities if hasattr(document, 'entities') else []

        parties: List[str] = []
        dates: List[str] = []
        contract_value: Any = None
        jurisdiction: Any = None
        governing_law: Any = None

        # 1) Prefer Document AI entities when present
        for entity in entities:
            entity_type = getattr(entity, 'type_', '')
            entity_value = getattr(entity, 'mention_text', '')
            if not entity_value:
                continue
            if entity_type in ['PERSON', 'ORGANIZATION']:
                parties.append(entity_value)
            elif entity_type in ['DATE', 'DUE_DATE', 'EFFECTIVE_DATE']:
                dates.append(entity_value)
            elif entity_type == 'CURRENCY':
                contract_value = contract_value or entity_value
            elif entity_type == 'LOCATION':
                jurisdiction = jurisdiction or entity_value
            elif entity_type == 'LAW':
                governing_law = governing_law or entity_value

        # 2) Fallback heuristic extraction from raw text
        if not parties or not dates or not governing_law or not contract_value:
            t_parties, t_dates, t_value, t_gov_law = self._parse_metadata_from_text(text)
            if not parties:
                parties = t_parties
            if not dates:
                dates = t_dates
            if not contract_value:
                contract_value = t_value
            if not governing_law:
                governing_law = t_gov_law

        document_type = self._classify_document_type(text)

        return DocumentMetadata(
            document_type=document_type,
            parties=list(dict.fromkeys([p.strip() for p in parties if p.strip()])),
            dates=list(dict.fromkeys([d.strip() for d in dates if d.strip()])),
            contract_value=contract_value,
            jurisdiction=jurisdiction,
            governing_law=governing_law
        )
    
    def _extract_clauses(self, document: Any) -> List[Clause]:
        """Extract individual clauses from the document"""
        clauses = []
        
        # Use Document AI's table and form extraction capabilities
        # This is a simplified implementation - in practice, you'd use more sophisticated parsing
        
        # Split document into sections based on common legal document patterns
        text = document.text
        sections = self._split_into_clauses(text)
        
        for i, section in enumerate(sections):
            if len(section.strip()) > 50:  # Only include substantial sections
                clause = Clause(
                    id=f"clause_{i+1}",
                    title=f"Section {i+1}",
                    text=section.strip(),
                    clause_type=self._classify_clause_type(section),
                    metadata={
                        "length": len(section),
                        "page_start": self._find_page_number(section, document),
                        "page_end": self._find_page_number(section, document)
                    },
                    start_page=self._find_page_number(section, document),
                    end_page=self._find_page_number(section, document)
                )
                clauses.append(clause)
        
        return clauses
    
    def _split_into_clauses(self, text: str) -> List[str]:
        """Split document text into logical clauses using robust legal heading detection."""
        if not text:
            return []

        # Normalize dashes
        norm = text.replace('–', '-').replace('—', '-')

        # Headings to detect: ARTICLE N - Title, Section N.N, numbered headings, ALL CAPS lines
        heading_regex = re.compile(
            r'(^\s*ARTICLE\s+\d+[^\n]*$)'
            r'|(^\s*Section\s+\d+(?:\.\d+)*[^\n]*$)'
            r'|(^\s*\d+\.\s+[^\n]+$)'
            r'|(^\s*[A-Z][A-Z0-9 \-/&]{3,}$)',
            re.MULTILINE
        )

        indices = [m.start() for m in heading_regex.finditer(norm)]
        # Ensure start included
        if 0 not in indices:
            indices = [0] + indices
        # Unique and sorted
        indices = sorted(set(indices))

        sections: List[str] = []
        for i, start in enumerate(indices):
            end = indices[i + 1] if i + 1 < len(indices) else len(norm)
            chunk = norm[start:end].strip()
            if len(chunk) > 50:
                sections.append(chunk)

        # Fallback to previous simple split if nothing found
        if not sections:
            section_patterns = [
                r'\n\s*\d+\.\s+',
                r'\n\s*[A-Z][A-Z\s]+\n',
                r'\n\s*Article\s+\d+',
                r'\n\s*Section\s+\d+',
                r'\n\s*Clause\s+\d+',
            ]
            sections = [norm]
            for pattern in section_patterns:
                new_sections = []
                for section in sections:
                    new_sections.extend(re.split(pattern, section))
                sections = new_sections
            sections = [s for s in sections if len(s.strip()) > 50]

        return sections

    def _parse_metadata_from_text(self, text: str) -> (List[str], List[str], Any, Any):
        """Heuristic extraction of parties, dates, contract value, and governing law from raw text."""
        parties: List[str] = []
        dates: List[str] = []
        contract_value: Any = None
        governing_law: Any = None

        if not text:
            return parties, dates, contract_value, governing_law

        # Parties line e.g., "Parties: Alpha Technologies, Inc. ("Provider") and Beta Retail LLC ("Client")."
        m = re.search(r'(?i)\bParties?:\s*(.+?)\.', text)
        if m:
            span = m.group(1)
            # Split around ' and ' while keeping commas inside names
            candidates = [p.strip() for p in re.split(r'\band\b', span) if p.strip()]
            # Remove role parentheses
            for c in candidates:
                name = re.sub(r'\s*\(.*?\)', '', c).strip(' "')
                if name:
                    parties.append(name)

        # Also look for Provider/Client named definitions
        for label in ["Provider", "Consultant", "Vendor", "Supplier", "Client", "Customer"]:
            m2 = re.search(rf'(?i)\b{label}\b[^\n]*?:\s*([A-Z][^\(\n\r]+)', text)
            if m2:
                name = m2.group(1).strip().strip(' "')
                if name and len(name.split()) <= 8:
                    parties.append(name)

        # Dates: capture formats like March 15, 2024 or 15 March 2024
        date_patterns = [
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s*\d{4}',
            r'\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December),?\s*\d{4}',
            r'\b\d{4}-\d{2}-\d{2}\b',
        ]
        for pat in date_patterns:
            for dm in re.findall(pat, text):
                if isinstance(dm, tuple):
                    match_text = re.search(pat, text)
                    if match_text:
                        dates.append(match_text.group(0))
                else:
                    dates.append(dm if isinstance(dm, str) and re.search(dm, text) is None else dm)
        # Also capture Effective Date line
        m3 = re.search(r'(?i)Effective Date:\s*([^\n\r]+)', text)
        if m3:
            dates.append(m3.group(1).strip())

        # Contract value: $ amounts
        m4 = re.search(r'\$\s?([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]{2})?)', text)
        if m4:
            contract_value = f"${m4.group(1)}"

        # Governing Law
        m5 = re.search(r'(?i)Governing Law:\s*([^\n\r]+)', text)
        if m5:
            governing_law = m5.group(1).strip()

        return parties, dates, contract_value, governing_law
    
    def _classify_clause_type(self, text: str) -> str:
        """Classify the type of clause based on content"""
        text_lower = text.lower()
        
        if any(keyword in text_lower for keyword in ['payment', 'compensation', 'salary', 'fee']):
            return 'payment'
        elif any(keyword in text_lower for keyword in ['termination', 'expire', 'end']):
            return 'termination'
        elif any(keyword in text_lower for keyword in ['liability', 'damages', 'indemnify']):
            return 'liability'
        elif any(keyword in text_lower for keyword in ['confidential', 'proprietary', 'secret']):
            return 'confidentiality'
        elif any(keyword in text_lower for keyword in ['governing law', 'jurisdiction', 'legal']):
            return 'governing_law'
        else:
            return 'general'
    
    def _classify_document_type(self, text: str) -> str:
        """Classify the type of legal document"""
        text_lower = text.lower()
        
        if any(keyword in text_lower for keyword in ['employment agreement', 'employment contract']):
            return 'employment_agreement'
        elif any(keyword in text_lower for keyword in ['service agreement', 'service contract']):
            return 'service_agreement'
        elif any(keyword in text_lower for keyword in ['lease agreement', 'rental agreement']):
            return 'lease_agreement'
        elif any(keyword in text_lower for keyword in ['purchase agreement', 'sales agreement']):
            return 'purchase_agreement'
        else:
            return 'general_contract'
    
    def _find_page_number(self, text: str, document: Any) -> int:
        """Find the page number where text appears (simplified)"""
        # This is a simplified implementation
        # In practice, you'd use Document AI's page information
        return 1
    
    def _calculate_confidence(self, document: Any) -> float:
        """Calculate confidence score for the document analysis"""
        # This is a simplified implementation
        # In practice, you'd analyze the quality of extraction
        return 0.85  # Default confidence score
    
    def _handle_processing_error(self, error: Exception, filename: str, session_id: str) -> None:
        """Handle document processing errors"""
        print(f"Error processing document {filename}: {str(error)}")
        raise error
